import os, re, sys, json
from docx import Document
from docx.oxml.ns import qn
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, Cm, Inches, RGBColor


def autoParagraph(document, paragraph_config):
    paragraph_alignment = {
        "CENTER": WD_ALIGN_PARAGRAPH.CENTER,
        "LEFT": WD_ALIGN_PARAGRAPH.LEFT,
        "RIGHT": WD_ALIGN_PARAGRAPH.RIGHT,
    }
    for paragraph in document.paragraphs:
        print(paragraph.style.name)
        if 'graphicData' in paragraph._p.xml:
            paragraph.paragraph_format.first_line_indent = 0
            paragraph.paragraph_format.left_indent = 0
            paragraph.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
            continue
        if paragraph.style.name in paragraph_config:
            key = paragraph.style.name
        else:
            # 修正 paragraph.style.name
            paragraph.style.name = "Normal"
        
        if paragraph.style.name == "Normal":
            if re.match("^图\d\s{1}", paragraph.text):
                key = "图"
            elif re.match("^表\d\s{1}", paragraph.text):
                key = "表"
            else:
                key = "Normal"
        
        paragraph_dict = paragraph_config[key]["paragraph_format"]
        font_dict = paragraph_config[key]["font_format"]
        paragraph.paragraph_format.alignment = paragraph_alignment[
            paragraph_dict["alignment"]
        ]
        # 首行悬挂缩进
        paragraph.paragraph_format.first_line_indent = Pt(
            paragraph_dict["first_line_indent"]
        )
        # 段落整体缩进
        paragraph.paragraph_format.left_indent = Pt(paragraph_dict["left_indent"])
        # 行距
        paragraph.paragraph_format.line_spacing = paragraph_dict["line_spacing"]
        # 前后间隔
        paragraph.paragraph_format.space_before = Pt(paragraph_dict["space_before"])
        paragraph.paragraph_format.space_after = Pt(paragraph_dict["space_after"])

        for run in paragraph.runs:
            # 大小
            run.font.size = Pt(font_dict["size"])
            # 字体
            if font_dict["name"] == "宋体":
                run.font.name = "宋体"
                run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
            elif font_dict["name"] == "黑体":
                run.font.name = "黑体"
                run._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")
            # 倾斜
            run.font.italic = font_dict["italic"]
            # 加粗
            run.font.bold = font_dict["bold"]
            # 下划线
            run.font.underline = font_dict["underline"]
            # 颜色
            if len(font_dict["color"]) != 8:
                print(key + " 颜色设置错误！")
                exit()
            hex_int = int(font_dict["color"], 16)
            if hex_int > 0xFFFFFF:
                print(key + " 颜色设置错误！")
                exit()

            b = hex_int & 0x0000FF
            g = (hex_int & 0x00FF00) >> 8
            r = (hex_int & 0xFF0000) >> 16
            run.font.color.rgb = RGBColor(r, g, b)


def autoTable(document, table_config):
    font_dict = table_config["font_format"]
    tables=document.tables
    b=0
    for tb in document.tables[b:]:
        for row in tb.rows:            
            for cell in row.cells:
                for par in cell.paragraphs:
                    for run in par.runs:                    
                        run.font.size = Pt(font_dict["size"])
                        if font_dict["name"] == "宋体":
                            run.font.name ='宋体'
                            run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
                        elif font_dict["name"] == "黑体":
                            run.font.name = "黑体"
                            run._element.rPr.rFonts.set(qn("w:eastAsia"), "黑体")
                        # 倾斜
                        run.font.italic = font_dict["italic"]
                        # 加粗
                        run.font.bold = font_dict["bold"]
                        # 下划线
                        run.font.underline = font_dict["underline"]
                        # 颜色
                        if len(font_dict["color"]) != 8:
                            print("table.json 颜色设置错误！")
                            exit()
                        hex_int = int(font_dict["color"], 16)
                        if hex_int > 0xFFFFFF:
                            print("table.json 颜色设置错误！")
                            exit()

                        b = hex_int & 0x0000FF
                        g = (hex_int & 0x00FF00) >> 8
                        r = (hex_int & 0xFF0000) >> 16
                        run.font.color.rgb = RGBColor(r, g, b)

if __name__ == "__main__":
    if len(sys.argv) <= 1:
        print("请输入源文件")
        exit(1)
    srcPath = sys.argv[1]

    dstPath = None
    if len(sys.argv) >= 3:
        dstPath = sys.argv[2]
    if dstPath is None:
        dstPath = srcPath
    if not os.path.exists(srcPath):
        print("源文件不存在！")
        exit(1)

    # srcPath = "./src/修改后的文件.docx"
    # dstPath = "./src/修改后的文件.docx"

    document = Document(srcPath)
    paragraph_config = None
    with open("/Users/book/Workspace/Python/autodoc/configs/paragraph.json") as f:
        paragraph_config = json.load(f)
    print(paragraph_config)
    with open("/Users/book/Workspace/Python/autodoc/configs/table.json") as f:
        table_config = json.load(f)
    autoParagraph(document, paragraph_config)
    autoTable(document, table_config)
    document.save(dstPath)
    print("自动格式刷完成 " + dstPath)
